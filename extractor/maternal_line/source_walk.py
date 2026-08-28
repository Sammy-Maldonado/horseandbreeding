"""Walk a .docx body into an ordered list of source nodes with stable identity.

The walk is the only place that reads through python-docx. It decides,
explicitly and testably, which body blocks are *meaningful* (they must be
accounted for downstream) and which are *structural* (section properties,
empty paragraphs) and therefore excluded from the ledger.
"""

from __future__ import annotations

import hashlib
import io
import os
from dataclasses import dataclass, field
from typing import Optional, Union

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run

Source = Union[str, "os.PathLike[str]", bytes, bytearray]


def normalise_text(text: str) -> str:
    return " ".join(text.split())


@dataclass
class SourceNode:
    node_id: str
    block_index: int
    kind: str  # paragraph | table | section_properties | other
    raw_text: str
    text: str
    meaningful: bool
    structural_reason: Optional[str] = None
    left_indent_pt: Optional[float] = None
    leading_tabs: int = 0
    bold_spans: list = field(default_factory=list)
    all_bold: bool = False
    cells: Optional[list] = None
    style_name: Optional[str] = None


@dataclass
class SourceWalk:
    document_id: str
    fingerprint: str
    nodes: list

    @property
    def meaningful_nodes(self) -> list:
        return [node for node in self.nodes if node.meaningful]

    @property
    def structural_nodes(self) -> list:
        return [node for node in self.nodes if not node.meaningful]


def load_source(source: Source, document_id: Optional[str] = None) -> tuple:
    if isinstance(source, (bytes, bytearray)):
        return bytes(source), document_id or "<bytes>"
    path = os.fspath(source)
    with open(path, "rb") as handle:
        data = handle.read()
    return data, document_id or os.path.basename(path)


def walk_source(source: Source, document_id: Optional[str] = None) -> SourceWalk:
    data, resolved_id = load_source(source, document_id)
    fingerprint = hashlib.sha256(data).hexdigest()
    document = docx.Document(io.BytesIO(data))
    prefix = fingerprint[:12]
    nodes = [
        _build_node(child, index, prefix, document)
        for index, child in enumerate(document.element.body.iterchildren())
    ]
    return SourceWalk(document_id=resolved_id, fingerprint=fingerprint, nodes=nodes)


def _build_node(child, index: int, prefix: str, document) -> SourceNode:
    node_id = f"{prefix}:b{index}"
    tag = child.tag
    if tag == qn("w:p"):
        return _paragraph_node(Paragraph(child, document._body), node_id, index)
    if tag == qn("w:tbl"):
        return _table_node(Table(child, document._body), node_id, index)
    if tag == qn("w:sectPr"):
        return SourceNode(node_id, index, "section_properties", "", "", False, "section_properties")
    raw = "".join(child.itertext())
    text = normalise_text(raw)
    return SourceNode(
        node_id, index, "other", raw, text,
        meaningful=bool(text),
        structural_reason=None if text else "empty_other",
    )


def _paragraph_node(paragraph: Paragraph, node_id: str, index: int) -> SourceNode:
    raw = paragraph.text
    text = normalise_text(raw)
    leading = raw[: len(raw) - len(raw.lstrip())]
    bold_spans, all_bold = _bold_evidence(paragraph)
    style = paragraph.style
    return SourceNode(
        node_id=node_id,
        block_index=index,
        kind="paragraph",
        raw_text=raw,
        text=text,
        meaningful=bool(text),
        structural_reason=None if text else "empty_paragraph",
        left_indent_pt=_left_indent_pt(paragraph),
        leading_tabs=leading.count("\t"),
        bold_spans=bold_spans,
        all_bold=all_bold,
        style_name=style.name if style is not None else None,
    )


def _left_indent_pt(paragraph: Paragraph) -> Optional[float]:
    indent = paragraph.paragraph_format.left_indent
    if indent is None and paragraph.style is not None:
        indent = paragraph.style.paragraph_format.left_indent
    if indent is None:
        return None
    return round(indent.pt, 2)


def _bold_evidence(paragraph: Paragraph) -> tuple:
    runs = []
    for content in paragraph.iter_inner_content():
        if isinstance(content, Run):
            runs.append(content)
        elif isinstance(content, Hyperlink):
            runs.extend(content.runs)
    style_bold = _style_bold(paragraph)
    spans: list = []
    bold_flags: list = []
    current = ""
    for run in runs:
        if not run.text:
            continue
        is_bold = run.bold if run.bold is not None else style_bold
        bold_flags.append(bool(is_bold))
        if is_bold:
            current += run.text
        elif current:
            spans.append(current)
            current = ""
    if current:
        spans.append(current)
    all_bold = bool(bold_flags) and all(bold_flags)
    return spans, all_bold


def _style_bold(paragraph: Paragraph) -> bool:
    style = paragraph.style
    while style is not None:
        if style.font is not None and style.font.bold is not None:
            return bool(style.font.bold)
        style = style.base_style
    return False


def _table_node(table: Table, node_id: str, index: int) -> SourceNode:
    cells = _table_cells(table)
    raw = "\n".join("\t".join(row) for row in cells)
    text = normalise_text(
        " || ".join(" | ".join(cell for cell in row if cell) for row in cells if any(row))
    )
    return SourceNode(
        node_id=node_id,
        block_index=index,
        kind="table",
        raw_text=raw,
        text=text,
        meaningful=True,  # a table is always a lot boundary, even when empty
        cells=cells,
    )


def _table_cells(table: Table) -> list:
    rows = []
    for row in table.rows:
        seen = set()
        texts = []
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:  # merged cells are repeated by python-docx
                continue
            seen.add(key)
            texts.append(_cell_text(cell))
        rows.append(texts)
    return rows


def _cell_text(cell) -> str:
    parts = [normalise_text(p.text) for p in cell.paragraphs if normalise_text(p.text)]
    for nested in cell.tables:
        for row in _table_cells(nested):
            parts.extend(text for text in row if text)
    return " ".join(parts)
