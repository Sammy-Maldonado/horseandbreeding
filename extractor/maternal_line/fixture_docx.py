"""
Programmatic .docx fixture builder for the extractor tests.

Every fixture is invented: no private catalogue text, no real horse, rider or
place names. Documents are built in memory with python-docx so the structural
intent (tables, paragraph order, bold runs, indentation, tabs) stays reviewable
in the test itself.

Block forms accepted by ``build_docx``:

    ("table", [["cell", "cell"], ["cell", "cell"]])
    ("p", "paragraph text")
    ("p", "paragraph text", {"indent": 14, "bold": ["1.40m"], "all_bold": True})

``indent`` is the paragraph left indentation in points. ``bold`` lists the
substrings that must be emitted as bold runs; everything else is a plain run.
"""

from __future__ import annotations

import io
import os
import tempfile

import docx
from docx.shared import Pt


def build_document(blocks):
    document = docx.Document()
    for block in blocks:
        kind = block[0]
        if kind == "table":
            _add_table(document, block[1])
        elif kind == "p":
            options = block[2] if len(block) > 2 else {}
            _add_paragraph(document, block[1], options)
        else:
            raise ValueError(f"unknown fixture block kind: {kind!r}")
    return document


def build_docx_bytes(blocks) -> bytes:
    buffer = io.BytesIO()
    build_document(blocks).save(buffer)
    return buffer.getvalue()


def write_docx(blocks, directory: str | None = None, name: str = "fixture.docx") -> str:
    """Write the fixture to disk (needed by the CLI tests) and return its path."""
    directory = directory or tempfile.mkdtemp(prefix="hb-extractor-")
    path = os.path.join(directory, name)
    with open(path, "wb") as handle:
        handle.write(build_docx_bytes(blocks))
    return path


def _add_table(document, rows):
    if not rows:
        table = document.add_table(rows=1, cols=1)
        return table
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    for r, row in enumerate(rows):
        for c, text in enumerate(row):
            table.cell(r, c).text = text
    return table


def _add_paragraph(document, text, options):
    paragraph = document.add_paragraph()
    bold_spans = list(options.get("bold", []))
    if options.get("all_bold"):
        run = paragraph.add_run(text)
        run.bold = True
    elif bold_spans:
        _add_runs_with_bold(paragraph, text, bold_spans)
    else:
        paragraph.add_run(text)
    indent = options.get("indent")
    if indent is not None:
        paragraph.paragraph_format.left_indent = Pt(indent)
    return paragraph


def _add_runs_with_bold(paragraph, text, bold_spans):
    """Emit ``text`` as runs, bolding the first occurrence of each listed span."""
    pending = list(bold_spans)
    cursor = 0
    while cursor < len(text):
        next_start, next_span = None, None
        for span in pending:
            found = text.find(span, cursor)
            if found != -1 and (next_start is None or found < next_start):
                next_start, next_span = found, span
        if next_start is None:
            paragraph.add_run(text[cursor:])
            break
        if next_start > cursor:
            paragraph.add_run(text[cursor:next_start])
        run = paragraph.add_run(next_span)
        run.bold = True
        pending.remove(next_span)
        cursor = next_start + len(next_span)
